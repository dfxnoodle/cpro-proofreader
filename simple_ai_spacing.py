#!/usr/bin/env python3
"""
Simple AI-based DOCX text spacing correction
Extracts text from DOCX and uses OpenAI to fix spacing issues
"""

import os
import sys
from pathlib import Path
from typing import Optional
from docx import Document
from openai import AzureOpenAI
from dotenv import load_dotenv
import argparse

# Load environment variables
load_dotenv()

class SimpleSpacingCorrector:
    """Simple text spacing corrector using AI"""
    
    def __init__(self):
        """Initialize with OpenAI client"""
        self.client = AzureOpenAI(
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            api_version="2025-01-01-preview"
        )
        
        if not os.getenv("AZURE_OPENAI_ENDPOINT") or not os.getenv("AZURE_OPENAI_API_KEY"):
            raise ValueError("Please set AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_API_KEY environment variables")
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract all text from a DOCX file"""
        try:
            doc = Document(docx_path)
            full_text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Skip empty paragraphs
                    full_text.append(paragraph.text)
            
            return '\n'.join(full_text)
        
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
    
    def correct_spacing_with_ai(self, text: str) -> str:
        """Send text to AI for spacing correction"""
        
        prompt = """You are a professional text editor. Please fix the spacing issues in the following text. 

Common issues to fix:
- Missing spaces between words that are stuck together
- Missing spaces after punctuation marks
- Extra spaces that shouldn't be there
- Proper spacing around Chinese/English mixed text
- Proper paragraph spacing

Please return only the corrected text without any explanations or additional comments.

Text to fix:
"""
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",  # Use your deployed model name
                messages=[
                    {
                        "role": "system", 
                        "content": "You are a professional text editor specializing in fixing spacing issues in documents. Return only the corrected text."
                    },
                    {
                        "role": "user", 
                        "content": prompt + text
                    }
                ],
                temperature=0.1,  # Low temperature for consistent corrections
                max_tokens=4000
            )
            
            return response.choices[0].message.content.strip()
        
        except Exception as e:
            raise Exception(f"Error calling OpenAI API: {str(e)}")
    
    def save_corrected_text(self, corrected_text: str, output_path: str):
        """Save corrected text to a new DOCX file"""
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading('Corrected Text', 0)
            
            # Add corrected text as paragraphs
            paragraphs = corrected_text.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text)
            
            doc.save(output_path)
            print(f"✅ Corrected text saved to: {output_path}")
        
        except Exception as e:
            raise Exception(f"Error saving corrected DOCX: {str(e)}")
    
    def process_file(self, input_path: str, output_path: Optional[str] = None) -> str:
        """Process a DOCX file: extract text, correct spacing, save result"""
        
        # Validate input file
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file not found: {input_path}")
        
        if not input_path.lower().endswith('.docx'):
            raise ValueError("Input file must be a .docx file")
        
        # Generate output path if not provided
        if output_path is None:
            input_stem = Path(input_path).stem
            output_path = f"{input_stem}_corrected.docx"
        
        print(f"📄 Processing: {input_path}")
        
        # Step 1: Extract text
        print("🔍 Extracting text from DOCX...")
        original_text = self.extract_text_from_docx(input_path)
        
        if not original_text.strip():
            raise ValueError("No text found in the DOCX file")
        
        print(f"📝 Extracted {len(original_text)} characters")
        
        # Step 2: Correct spacing with AI
        print("🤖 Sending to AI for spacing correction...")
        corrected_text = self.correct_spacing_with_ai(original_text)
        
        print(f"✨ Received {len(corrected_text)} characters back")
        
        # Step 3: Save corrected text
        print("💾 Saving corrected text...")
        self.save_corrected_text(corrected_text, output_path)
        
        # Show comparison
        print("\n" + "="*60)
        print("ORIGINAL TEXT (first 200 chars):")
        print("="*60)
        print(original_text[:200] + "..." if len(original_text) > 200 else original_text)
        
        print("\n" + "="*60)
        print("CORRECTED TEXT (first 200 chars):")
        print("="*60)
        print(corrected_text[:200] + "..." if len(corrected_text) > 200 else corrected_text)
        print("="*60)
        
        return output_path


def main():
    """Command line interface"""
    parser = argparse.ArgumentParser(description='Fix spacing in DOCX files using AI')
    parser.add_argument('input_file', help='Input DOCX file path')
    parser.add_argument('-o', '--output', help='Output DOCX file path (optional)')
    
    args = parser.parse_args()
    
    try:
        corrector = SimpleSpacingCorrector()
        output_file = corrector.process_file(args.input_file, args.output)
        
        print(f"\n🎉 Success! Corrected file saved as: {output_file}")
        
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()
