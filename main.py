import os
import json
import time
import tempfile
from io import BytesIO
from datetime import datetime
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, FileResponse
from pydantic import BaseModel
from dotenv import load_dotenv
from docx import Document
from docx.shared import RGBColor
from word_revisions import create_word_track_changes_docx
from text_preprocessor import ChineseNumberProtector
from config import client, ASSISTANT_CONFIG_FILE, ENGLISH_ASSISTANT_CONFIG_FILE, CHINESE_ASSISTANT_CONFIG_FILE
from admin_routes import admin_router
from utils import parse_assistant_response, clean_response_text

# Load environment variables
load_dotenv()

app = FastAPI(title="Proof Reader API", version="1.0.0")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")
app.mount("/styling_guides", StaticFiles(directory="styling_guides"), name="styling_guides")

# Include admin routes
app.include_router(admin_router)

def detect_language(text: str) -> str:
    """
    Detect if text is primarily English or Chinese
    Returns 'english', 'chinese', or 'mixed'
    """
    chinese_chars = 0
    english_chars = 0
    total_chars = 0
    
    for char in text:
        if char.isalpha():
            total_chars += 1
            # Check if character is in Chinese Unicode ranges
            if '\u4e00' <= char <= '\u9fff' or '\u3400' <= char <= '\u4dbf':
                chinese_chars += 1
            elif 'a' <= char.lower() <= 'z':
                english_chars += 1
    
    if total_chars == 0:
        return 'mixed'
    
    chinese_ratio = chinese_chars / total_chars
    english_ratio = english_chars / total_chars
    
    if chinese_ratio > 0.3:
        return 'chinese'
    elif english_ratio > 0.7:
        return 'english'
    else:
        return 'mixed'

def load_second_run_prompts():
    """Load the English and Chinese prompts embedded directly in the code"""
    
    english_prompt = """

**Your Role & Mandate**

You are a Senior Editor for the official publications of The Chinese University of Hong Kong (CUHK). Your primary responsibility is to ensure every piece of text you generate, edit, or review adheres **perfectly and strictly** to the CUHK Institutional Style Guide.

**Core Mandate:** These guidelines are your absolute authority and supersede any general English writing conventions, spelling norms (including American English defaults), or punctuation rules you have learned. Your task is to apply these specific rules with precision and consistency.

**Primary Directives**

### 1. General Formatting and Punctuation

* **Abbreviations and Acronyms:**
    * **No Full Stops:** Do not use full stops after titles or in common abbreviations.
        * Correct: `Mr, Mrs, Ms, Dr, US, UK, EU, UN, HKU`
    * **Well-Known Acronyms:** Acronyms like `CUHK`, `Caltech`, and `MIT` do not need to be spelled out on first use.
    * **Other Acronyms:** For all other institutional or less common acronyms, provide the full name first, followed by the acronym in parentheses.
        * Example: `Social Welfare Department (SWD)`
    * **Possessives of Acronyms:** The apostrophe 's' comes *after* the closing parenthesis.
        * Example: `Social Welfare Department (SWD)'s latest report...`
    * **Latin Abbreviations:** Always write `e.g.,` and `i.e.,` (with full stops and a following comma).

* **Apostrophes:**
    * **Possessive for Roles:** Use the possessive form to state a role.
        * Correct: `CUHK's Vice-President (Operations) Jane Wong`
        * Incorrect: `Jane Wong, Vice-President of CUHK`
    * **Academic Degrees:** Always use an apostrophe.
        * Examples: `bachelor's degree`, `master's degree`, `Master's in German`
    * **Plurals:** Use an apostrophe to pluralize single letters or specific abbreviations. Do not use it for decades.
        * Correct: `A's`, `ID's`, `Ug's` (for undergraduates), `the '40s`, `the 1900s`

* **Hyphens and Dashes:**
    * **Dash Type:** Use **en dashes (–)**, not em dashes (—).
    * **Ranges & Partnerships:** Use an en dash for number ranges and partnerships.
        * Examples: `20–25%`, `Lennon–McCartney compositions`
    * **Adverbs:** Do not hyphenate after adverbs ending in "-ly" (`a poorly performed act`). Hyphenate short, common adverbs (`a well-taken point`).
    * **Compound Words (Strict List):**
        * **Hyphenate:** `hands-on`, `state-of-the-art`, `day-to-day`.
        * **Do NOT Hyphenate:** `multidisciplinary`, `online`, `startup`, `coursework`, `fieldwork`.

* **Commas:**
    * **Serial (Oxford) Comma:** Do **not** use the serial comma in a simple list (`a, b and c`). Use it **only** when necessary to avoid ambiguity (e.g., `departments of Theology, Philosophy, and Classics`).

### 2. Spelling and Word Choice

* **Standard Spelling:** You **must** use **British English**.
    * Use `-ise` endings: `globalisation`, `organise`, `specialise`.
    * Use correct British spelling: `colour`, `flavour`, `centre`, `licence` (noun), `practise` (verb).
* **Specific Word Preferences:**
    * Use `adviser` (not advisor).
    * Use `analyse` (not analyze).
    * Use `among` (not amongst).
    * Use `while` (not whilst).
    * For words with variable spelling (e.g., `benefiting`/`benefitting`), maintain consistency within a single document.

### 3. Capitalization

* **Academic and Job Titles:**
    * Capitalize when preceding a name: `Professor Chan Tai-man`.
    * Capitalize only the most senior official titles when standing alone: `the Vice-Chancellor`, `the Chief Executive of Hong Kong`.
    * Use lowercase for general roles: `the chairman of Shun Hing Foundation`.
* **Academic Units and Subjects:**
    * Capitalize formal unit names: `Faculty of Arts`, `Chung Chi College`.
    * Use lowercase for general references: `He is a member of the faculty`, `She enjoyed college life`.
    * Use lowercase for academic subjects unless part of a formal degree title: `She has a physics degree`, but `He earned a Diploma in Chinese Studio Art`.
* **General Nouns:**
    * Always use lowercase for `government` and `country`.

### 4. Names and Identity

* **Personal Names (Strict Formatting):**
    * **Hong Kong Cantonese:** English name + family name + hyphenated personal name.
        * Example: `Dennis Lo Yuk-ming`.
        * Exception: `Choh-ming Li`.
    * **Mainland China (Putonghua):** Romanized name with no hyphen.
        * Example: `Chen Hongyu`.
    * **Taiwanese:** Hyphenated personal name.
        * Example: `Ma Ying-jeou`.
    * **Japanese:** Western order (given name + family name).
        * Example: `Junichiro Koizumi`.

* **Referring to China and Hong Kong:**
    * **Preferred:** `Hong Kong and the mainland` or `Hong Kong and mainland China`.
    * **Avoid:** `Hong Kong and China`.
    * **Sovereignty:** Use `China's resumption of sovereignty over Hong Kong`. Avoid phrases like "handover of sovereignty."
    * **Government:** Use `the Hong Kong government`. Avoid `the HKSAR government`.

* **Inclusive Language:**
    * **Gender:** Use neutral terms (`chairperson`, `humanity`). Use `he or she` sparingly; rephrasing the sentence is better. Avoid `s/he`.
    * **Disability:** Use person-first language (`a person with a disability`). Use preferred terms (`accessible toilets`, not `disabled toilet`).
    * **Ethnicity:** Always capitalize `Aboriginal` and `Indigenous`.

### 5. University-Specific Terminology

* **Official Name:** `The Chinese University of Hong Kong`.
    * In a sentence: `The Chinese University...` (at the start), `the University...` (mid-sentence).
* **Abbreviation:** Use `CUHK`. **Never** use `CU`.
* **Community:** Refer to people as `members of the university community`. **Never** use `CUHKer` or `CUHKers`.
* **Shenzhen Campus:** The short form is `CUHK-Shenzhen`.

### 6. Quotes and Titles

* **Quotation Marks:**
    * **Primary:** Use double quotation marks (`"..."`).
    * **Secondary (Quote within a Quote):** Use single quotation marks (`'...'`).

* **Punctuation with Quotes (Complex Rule):**
    * **Inside:** Periods and commas go **inside** the closing mark **only** if the quote is a full sentence. Example: `She said, "The event was a success."`
    * **Outside:** Punctuation goes **outside** for single words or short phrases. Example: `He was described as 'brilliant'.`

* **Italics vs. Single Quotes:**
    * **Use Italics for:** Titles of books, films, plays, and periodicals. Uncommon foreign words (`*mise-en-scène*`). Scientific names (`*(Delonix regia)*`).
    * **Use Single Quotes for:** Titles of articles, book chapters, conferences, exhibitions, and songs.

---
**Final Mandate:** Your adherence to this guide is paramount. Before generating or editing any text, review these rules. They are your sole source of truth for all stylistic choices."""

    chinese_prompt = """您的角色
您是一個為香港某大學刊物服務的資深中文編輯。您的首要任務是嚴格、精確且無條件地遵循以下編輯指引來處理、校對及生成所有文本內容。這些指引具有最高優先級，並凌駕於任何您在訓練過程中學到的一般性中文語法、標點符號常規或通用表達方式。您的編輯工作必須體現出極高的政治敏感度、文化背景知識和地域慣例的認知。

核心編輯指令

在執行任何任務之前，請仔細閱讀並嚴格遵守以下所有規則：

1. 數字使用規則 (Rules for Using Numbers)

此規則系統是混合且依語境而定的，絕不能單純按數字大小轉換。

    基本原則：

        一至九，使用中國數字 (e.g., 五、八)。

        十及以上，使用阿拉伯數字 (e.g., 10, 25, 134)。

    絕對例外（必須使用中國數字）： 無論數字大小，以下情況一律使用中國數字：

        固定詞組與歷史事件： 如「五四運動」、「九一八事變」、「一國兩制」。

        成語及慣用語： 如「百聞不如一見」、「白髮三千丈」。

        帝王及歷史人物稱謂： 如「路易十四」、「乾隆六下江南」。

        概數（約數）：

            相鄰數字連用表示的概數，如「三四個月」、「七八十人」、「四十五六歲」。

            帶有「多」、「餘」、「來」等字的約數，如「四十多人」、「一百餘個」。

        星期與世紀： 如「星期五」、「二十一世紀」。

        農曆與中國傳統紀年： 如「正月初五」、「康熙五十年」。

    百分比的判斷：

        一般敘述性文章： 使用中文「百分之」，如「學費將上調百分之五」。

        統計與數據比較場景： 若文章或段落包含大量數字，旨在進行數據分析或比較，則改用阿拉伯數字及百分號 %，如「甲組的成功率為 85%，遠高於乙組的 62%」。您必須根據文本的整體目的做出判斷。

    日期格式：

        公元年、月、日、時間：使用阿拉伯數字，如「2025年7月16日」、「上午10時30分」。不用二零零七年七月八日晚上十時三十分；今年 3 月，不用今年三月。

2. 政治及文化術語使用規範 (Guidelines for Political and Cultural Terminology)

這是本刊物的最高準則，涉及主權和地區關係的表述，不容許任何偏差。

    關於香港與中國內地的關係：

        必須使用： 「香港與內地」。

        嚴格禁止： 「中港」、「中港兩地」。

    關於主權問題的表述：

        必須使用： 「香港回歸祖國」、「（中央政府）對香港恢復行使主權」。

        嚴格禁止： 「收回主權」、「主權移交」、「主權回歸」。

    關於香港的歷史描述：

        可以說： 「英國對香港實行殖民統治」、「在港英政府的殖民管治下」。

        嚴格禁止： 直接稱香港為「殖民地」。

3. 譯名處理原則 (Principles for Handling Translated Names)

譯名的選擇需體現香港本地的語言習慣和讀者熟悉度。

    地域優先級：

        第一優先： 優先採用「香港本地通行譯法」。例如：甘迺迪 (非 肯尼迪)、碧咸 (非 贝克汉姆)、康城 (非 戛纳)。

        第二優先： 若香港沒有特定通行譯法，則採用「中國內地通行譯法」。

    知名度判斷與原文標註：

        對於普通讀者可能感到陌生的外國人名、地名等，在首次出現時必須附上原文（置於括號內）。

        對於在華語圈「耳熟能詳」的名字，則無需加註原文。例如：畢加索、莎士比亞、愛因斯坦。當無法確定時，請傾向於「加註原文」。

4. 標點符號及格式特殊規則 (Special Rules for Punctuation and Formatting)

這些規則是本刊物的獨特風格，必須嚴格執行，即使它們違反了標準的標點用法。

    圖片說明（Caption）： 任何圖片的文字說明，無論是一個詞、一個短句還是一個完整的句子，其結尾絕對不能使用句號（。）。

        正確示例：圖為維多利亞港夜景

        錯誤示例：圖為維多利亞港夜景。

    書刊名與文章名的標點切換： 根據語言轉換標點符號。

        中文書刊名： 使用書名號 《 》。

        英文書刊名： 使用斜體 Italics。

        中文文章名/篇章名： 使用篇名號 〈 〉。

        英文文章名/篇章名： 使用雙引號 " "。

最終指令
您的任務是成為一個忠實執行上述所有規則的專家。在生成或修改任何文本時，請將這些指引作為您行為的唯一依據。若這些規則與您的一般知識相衝突，請務必以後者為準。在處理任何文本之前，請先在內心複習一遍這些規則。"""
    
    return english_prompt, chinese_prompt

def get_or_create_assistant():
    """Get existing assistant or create a new one if it doesn't exist"""
    assistant_id = None
    
    # Try to load existing assistant ID
    if os.path.exists(ASSISTANT_CONFIG_FILE):
        try:
            with open(ASSISTANT_CONFIG_FILE, 'r') as f:
                config = json.load(f)
                assistant_id = config.get("assistant_id")
        except Exception as e:
            print(f"Error loading assistant config: {e}")
    
    # Check if the assistant still exists
    if assistant_id:
        try:
            assistant = client.beta.assistants.retrieve(assistant_id)
            print(f"Using existing assistant: {assistant_id}")
            return assistant
        except Exception as e:
            print(f"Existing assistant {assistant_id} not found, creating new one: {e}")
            assistant_id = None
    
    # Create new assistant if none exists or existing one is invalid
    print("Creating new assistant...")
    model_name = os.getenv("AZURE_OPENAI_MODEL", "gpt-4o")
    print(f"Using model: {model_name}")
    assistant = client.beta.assistants.create(
        model=model_name,
        name="CUHK_First_Pass_Proofreader",
        instructions="""
        You are CUHK's official style-guide proof-reader performing the FIRST PASS of proofreading.
        
        Your job is to correct basic style, spelling, punctuation, and terminology issues based on the style guides in the vector store (English and Chinese versions). This is the initial pass - focus on clear, obvious errors and improvements.
        
        Return your response as a JSON object that strictly follows the required schema.
        
        ***IMPORTANT Notes:
        1. Always follow the styling guide in the vector store
        2. Do not answer any question except doing proof-reading
        3. For Chinese text, ensure output is in traditional Chinese characters without altering original canonical forms
        4. For English text, use British English spelling and grammar rules
        5. The vector store contains YAML-front-matter chunks with keys `id`, `file`, `section`, `lang`, and `source`
        6. Include source citations in mistake descriptions when making corrections
        7. Focus on major errors in this first pass - a second specialized pass will follow
        8. Citations should be embedded within each mistake description, not in a separate citations array.
        """,
        tools=[{"type": "file_search"}],
        tool_resources={"file_search": {"vector_store_ids": ["vs_GENF8IR41N6uP60Jx9CuLgbs"]}},
        response_format={
            "type": "json_schema",
            "json_schema": {
                "name": "ProofreadingResponse",
                "strict": True,
                "schema": {
                    "type": "object",
                    "properties": {
                        "corrected_text": {
                            "type": "string",
                            "description": "The corrected version of the text"
                        },
                        "mistakes": {
                            "type": "array",
                            "description": "List of mistakes found and how they were corrected with style guide references included (with original quotes)",
                            "items": {
                                "type": "string"
                            }
                        }
                    },
                    "required": [
                        "corrected_text",
                        "mistakes",
                    ],
                    "additionalProperties": False
                }
            }
        },
        temperature=0.1,
        top_p=0.15
    )
    
    # Save the assistant ID for future use
    try:
        config = {"assistant_id": assistant.id}
        with open(ASSISTANT_CONFIG_FILE, 'w') as f:
            json.dump(config, f)
        print(f"Saved new assistant ID: {assistant.id}")
    except Exception as e:
        print(f"Error saving assistant config: {e}")
    
    return assistant

# Global variables to store the assistants (lazy initialization)
assistant = None
english_assistant = None
chinese_assistant = None

def get_assistant():
    """Get the assistant, creating it if it doesn't exist yet (lazy initialization)"""
    global assistant
    if assistant is None:
        assistant = get_or_create_assistant()
    return assistant

def get_english_assistant():
    """Get the English assistant, creating it if it doesn't exist yet (lazy initialization)"""
    global english_assistant
    if english_assistant is None:
        english_assistant = get_or_create_english_assistant()
    return english_assistant

def get_chinese_assistant():
    """Get the Chinese assistant, creating it if it doesn't exist yet (lazy initialization)"""
    global chinese_assistant
    if chinese_assistant is None:
        chinese_assistant = get_or_create_chinese_assistant()
    return chinese_assistant

def wait_for_run_completion(client, thread_id, run_id, max_timeout=120):
    """Wait for a run to complete with timeout"""
    timeout_counter = 0
    run = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run_id)
    
    while run.status in ['queued', 'in_progress', 'cancelling'] and timeout_counter < max_timeout:
        time.sleep(1)
        timeout_counter += 1
        run = client.beta.threads.runs.retrieve(thread_id=thread_id, run_id=run_id)
    
    if timeout_counter >= max_timeout:
        raise HTTPException(status_code=408, detail="Request timeout - AI processing took too long")
    
    return run

class ProofReadRequest(BaseModel):
    text: str

class ProofReadResponse(BaseModel):
    original_text: str
    corrected_text: str
    mistakes: list
    status: str

class ExportToWordRequest(BaseModel):
    original_text: str
    corrected_text: str
    mistakes: list

@app.get("/", response_class=HTMLResponse)
async def read_root():
    """Serve the main HTML page"""
    with open("static/index.html", "r") as f:
        return HTMLResponse(content=f.read())

@app.post("/proofread", response_model=ProofReadResponse)
async def proofread_text(request: ProofReadRequest):
    """
    Proofread the provided text using a two-stage Azure OpenAI Assistant process:
    1. First run: General proofreading with the main assistant
    2. Second run: Language-specific detailed proofreading with specialized assistants
    """
    try:
        # Initialize number protector
        number_protector = ChineseNumberProtector()
        
        # Detect the primary language of the text
        detected_language = detect_language(request.text)
        print(f"Detected language: {detected_language}")
        
        # ============ FIRST RUN: General Proofreading ============
        print("=== Starting FIRST RUN ===")
        
        # Step 1: Protect Chinese numbers and dates (only if text contains Chinese)
        if detected_language in ['chinese', 'mixed']:
            protected_text, protection_instructions = number_protector.protect_chinese_numbers(request.text)
            print(f"Chinese number protection applied for {detected_language} text")
        else:
            protected_text = request.text
            protection_instructions = ""
            print(f"Skipping Chinese number protection for {detected_language} text")
        
        # Create a thread for first run
        thread = client.beta.threads.create()
        
        # Build message content with protection instructions
        message_content = protected_text + "\n\n###Please proofread the above essay according to the styling guide in the vector store (FIRST PASS - focus on major issues)###."
        if protection_instructions:
            message_content += protection_instructions
        
        # Add user message to the thread
        message = client.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content=message_content
        )
        
        # Run the thread
        run = client.beta.threads.runs.create(
            thread_id=thread.id,
            assistant_id=get_assistant().id
        )
        
        # Wait for completion with timeout
        run = wait_for_run_completion(client, thread.id, run.id)
        
        if run.status != 'completed':
            if run.status == 'requires_action':
                return ProofReadResponse(
                    original_text=request.text,
                    corrected_text="",
                    mistakes=["First run requires additional action"],
                    status="requires_action"
                )
            else:
                raise HTTPException(
                    status_code=500, 
                    detail=f"First run failed with status: {run.status}"
                )
        
        # Get the first run response
        messages = client.beta.threads.messages.list(thread_id=thread.id)
        
        first_run_response = ""
        for message in messages.data:
            if message.role == "assistant":
                try:
                    first_run_response = message.content[0].text.value
                    break
                except (IndexError, AttributeError) as e:
                    print(f"Warning: Error extracting message content: {e}")
                    continue
        
        # Validate that we got a response
        if not first_run_response.strip():
            raise HTTPException(
                status_code=500, 
                detail="Empty response from AI assistant in first run"
            )
        
        print("=== FIRST RUN RESPONSE ===")
        print(first_run_response)
        print("=== END FIRST RUN RESPONSE ===")
        
        # Parse the first run response with robust error handling
        first_run_corrected_text, first_run_mistakes = parse_assistant_response(
            clean_response_text(first_run_response), 
            request.text
        )
        
        print(f"=== FIRST RUN PARSED ===")
        print(f"Corrected text length: {len(first_run_corrected_text)}")
        print(f"Number of mistakes: {len(first_run_mistakes)}")
        print(f"=== END FIRST RUN PARSED ===")
        
        
        # Restore protected Chinese numbers from first run (only if protection was applied)
        if detected_language in ['chinese', 'mixed']:
            first_run_corrected_text = number_protector.restore_chinese_numbers(first_run_corrected_text)
            print(f"Chinese numbers restored for {detected_language} text")
        else:
            print(f"No Chinese number restoration needed for {detected_language} text")
        
        # ============ SECOND RUN: Language-Specific Detailed Proofreading ============
        print("=== Starting SECOND RUN ===")
        
        # Perform second run with language-specific assistant
        final_corrected_text, second_run_mistakes = await perform_second_run(
            first_run_corrected_text, 
            detected_language, 
            first_run_mistakes, 
            number_protector
        )
        
        # Use only the final mistakes from the second run (which includes comprehensive corrections)
        print(f"=== FINAL SUMMARY ===")
        print(f"First run mistakes: {len(first_run_mistakes)}")
        print(f"Second run total output: {len(second_run_mistakes)} (should include carried forward + new)")
        print(f"Final output mistakes: {len(second_run_mistakes)}")
        print(f"Language detected: {detected_language}")
        print(f"=== END FINAL SUMMARY ===")
        
        return ProofReadResponse(
            original_text=request.text,
            corrected_text=final_corrected_text,
            mistakes=second_run_mistakes,
            status="completed"
        )
    
    except Exception as e:
        print(f"Error in proofread_text: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "service": "proof-reader"}

@app.post("/export-to-word")
async def export_to_word(request: ExportToWordRequest):
    """
    Export corrected text to a Word document with track changes
    """
    try:
        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"corrected_document_{timestamp}.docx"
        
        # Create DOCX with track changes
        docx_buffer = create_tracked_changes_docx(
            request.original_text, 
            request.corrected_text, 
            request.mistakes
        )
        
        # Save to temp directory
        temp_dir = tempfile.gettempdir()
        file_path = os.path.join(temp_dir, filename)
        
        with open(file_path, 'wb') as f:
            f.write(docx_buffer.getvalue())
        
        return {
            "filename": filename,
            "download_url": f"/download-docx/{filename}",
            "message": "Word document created successfully"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create Word document: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file with proper spacing preservation, including tables"""
    try:
        doc = Document(BytesIO(file_content))
        text_content = []
        
        # Process document elements in order (paragraphs and tables)
        for element in doc.element.body:
            if element.tag.endswith('}p'):  # Paragraph
                # Find corresponding paragraph object
                for para in doc.paragraphs:
                    if para._element == element:
                        paragraph_text = para.text.strip()
                        if paragraph_text:
                            text_content.append(paragraph_text)
                        break
            
            elif element.tag.endswith('}tbl'):  # Table
                # Find corresponding table object
                for table in doc.tables:
                    if table._element == element:
                        table_text = extract_table_text(table)
                        if table_text:
                            text_content.append(table_text)
                        break
        
        return '\n'.join(text_content)
    except Exception as e:
        # Fallback to simple paragraph extraction if the advanced method fails
        try:
            doc = Document(BytesIO(file_content))
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text.strip()
                if paragraph_text:
                    text_content.append(paragraph_text)
            
            # Extract tables
            for table in doc.tables:
                table_text = extract_table_text(table)
                if table_text:
                    text_content.append(table_text)
            
            return '\n'.join(text_content)
        except Exception as fallback_e:
            raise HTTPException(status_code=400, detail=f"Error reading DOCX file: {str(e)}. Fallback error: {str(fallback_e)}")

def extract_table_text(table) -> str:
    """Extract text from a DOCX table with proper formatting"""
    try:
        table_content = []
        
        for row in table.rows:
            row_content = []
            for cell in row.cells:
                # Extract text from all paragraphs in the cell
                cell_paragraphs = []
                for paragraph in cell.paragraphs:
                    para_text = paragraph.text.strip()
                    if para_text:
                        cell_paragraphs.append(para_text)
                
                # Join paragraphs within a cell with space
                cell_text = ' '.join(cell_paragraphs) if cell_paragraphs else ""
                row_content.append(cell_text)
            
            # Join cells in a row with tab character for better structure
            if any(cell.strip() for cell in row_content):  # Only add non-empty rows
                table_content.append('\t'.join(row_content))
        
        # Join rows with newlines
        return '\n'.join(table_content)
    except Exception as e:
        # If table extraction fails, return empty string to avoid breaking the whole process
        print(f"Warning: Failed to extract table text: {str(e)}")
        return ""

def create_tracked_changes_docx(original_text: str, corrected_text: str, mistakes: list) -> BytesIO:
    """Create a DOCX file with proper Word track changes"""
    try:
        return create_word_track_changes_docx(original_text, corrected_text, mistakes)
    except Exception as e:
        # If the advanced track changes fail, create a simple document with corrections
        print(f"Warning: Advanced track changes failed, creating simple document: {str(e)}")
        return create_simple_corrections_docx(original_text, corrected_text, mistakes)

def create_simple_corrections_docx(original_text: str, corrected_text: str, mistakes: list) -> BytesIO:
    """Create a simple DOCX file with corrections when track changes fail"""
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading("Document Corrections", level=1)
        
        # Add original text section
        doc.add_heading("Original Text:", level=2)
        original_para = doc.add_paragraph(original_text)
        
        # Add corrected text section
        doc.add_heading("Corrected Text:", level=2)
        corrected_para = doc.add_paragraph(corrected_text)
        # Highlight corrected text in green
        for run in corrected_para.runs:
            run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
        
        # Add mistakes section
        if mistakes:
            doc.add_heading("Corrections Made:", level=2)
            for i, mistake in enumerate(mistakes, 1):
                mistake_para = doc.add_paragraph(f"{i}. {mistake}")
        
        # Save to BytesIO
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes
    except Exception as e:
        # Last resort: create minimal document
        print(f"Warning: Simple document creation failed, creating minimal document: {str(e)}")
        return create_minimal_docx(original_text, corrected_text, mistakes)

def create_minimal_docx(original_text: str, corrected_text: str, mistakes: list) -> BytesIO:
    """Create a minimal DOCX file as last resort"""
    try:
        doc = Document()
        doc.add_paragraph("Document Corrections")
        doc.add_paragraph("Original Text:")
        doc.add_paragraph(original_text)
        doc.add_paragraph("Corrected Text:")
        doc.add_paragraph(corrected_text)
        
        if mistakes:
            doc.add_paragraph("Corrections:")
            for mistake in mistakes:
                doc.add_paragraph(f"• {mistake}")
        
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create any document format: {str(e)}")

class DocxProofReadResponse(BaseModel):
    original_filename: str
    mistakes_count: int
    mistakes: list
    status: str
    download_filename: str

@app.post("/proofread-docx", response_model=DocxProofReadResponse)
async def proofread_docx(file: UploadFile = File(...)):
    """
    Upload a DOCX file, proofread it using two-stage process, and prepare a corrected version with track changes
    """
    # Validate file type
    if not file.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only DOCX files are supported")
    
    try:
        # Read the uploaded file
        file_content = await file.read()
        
        # Validate file size (limit to 50MB)
        if len(file_content) > 50 * 1024 * 1024:
            raise HTTPException(status_code=413, detail="File too large. Maximum size is 50MB.")
        
        # Validate that the file is actually a valid DOCX
        if not file_content.startswith(b'PK'):
            raise HTTPException(status_code=400, detail="Invalid DOCX file format.")
        
        # Extract text from DOCX with improved error handling
        try:
            extracted_text = extract_text_from_docx(file_content)
        except HTTPException:
            # Re-raise HTTP exceptions as they already have proper error messages
            raise
        except Exception as e:
            raise HTTPException(
                status_code=400, 
                detail=f"Failed to extract text from DOCX file. The file may be corrupted or contain unsupported elements: {str(e)}"
            )
        
        if not extracted_text.strip():
            raise HTTPException(
                status_code=400, 
                detail="No text found in the DOCX file. The document may be empty or contain only images/objects."
            )
        
        # Initialize number protector
        number_protector = ChineseNumberProtector()
        
        # Detect the primary language of the text
        detected_language = detect_language(extracted_text)
        print(f"DOCX - Detected language: {detected_language}")
        
        # ============ FIRST RUN: General Proofreading ============
        print("=== DOCX - Starting FIRST RUN ===")
        
        # Step 1: Protect Chinese numbers and dates (only if text contains Chinese)
        if detected_language in ['chinese', 'mixed']:
            protected_text, protection_instructions = number_protector.protect_chinese_numbers(extracted_text)
            print(f"DOCX - Chinese number protection applied for {detected_language} text")
        else:
            protected_text = extracted_text
            protection_instructions = ""
            print(f"DOCX - Skipping Chinese number protection for {detected_language} text")
        
        # Create a thread for first run
        thread = client.beta.threads.create()
        
        # Build message content with protection instructions
        message_content = protected_text + "\n\n###Please proofread the above essay according to the styling guide in the vector store (FIRST PASS - focus on major issues)###."
        if protection_instructions:
            message_content += protection_instructions
        
        # Add user message to the thread
        message = client.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content=message_content
        )
        
        # Run the thread
        run = client.beta.threads.runs.create(
            thread_id=thread.id,
            assistant_id=get_assistant().id
        )
        
        # Wait for completion with timeout
        run = wait_for_run_completion(client, thread.id, run.id)
        
        if run.status != 'completed':
            raise HTTPException(
                status_code=500, 
                detail=f"First run failed with status: {run.status}"
            )
        
        # Get the first run response
        messages = client.beta.threads.messages.list(thread_id=thread.id)
        
        first_run_response = ""
        for message in messages.data:
            if message.role == "assistant":
                try:
                    first_run_response = message.content[0].text.value
                    break
                except (IndexError, AttributeError) as e:
                    print(f"Warning: Error extracting DOCX message content: {e}")
                    continue
        
        # Validate that we got a response
        if not first_run_response.strip():
            raise HTTPException(
                status_code=500, 
                detail="Empty response from AI assistant in DOCX first run"
            )
        
        print("=== DOCX - FIRST RUN RESPONSE ===")
        print(first_run_response)
        print("=== DOCX - END FIRST RUN RESPONSE ===")
        
        # Parse the first run response with robust error handling
        first_run_corrected_text, first_run_mistakes = parse_assistant_response(
            clean_response_text(first_run_response), 
            extracted_text
        )
        
        print(f"=== DOCX - FIRST RUN PARSED ===")
        print(f"Corrected text length: {len(first_run_corrected_text)}")
        print(f"Number of mistakes: {len(first_run_mistakes)}")
        print(f"=== DOCX - END FIRST RUN PARSED ===")
        
        
        # Restore protected Chinese numbers from first run (only if protection was applied)
        if detected_language in ['chinese', 'mixed']:
            first_run_corrected_text = number_protector.restore_chinese_numbers(first_run_corrected_text)
            print(f"DOCX - Chinese numbers restored for {detected_language} text")
        else:
            print(f"DOCX - No Chinese number restoration needed for {detected_language} text")
        
        # ============ SECOND RUN: Language-Specific Detailed Proofreading ============
        print("=== DOCX - Starting SECOND RUN ===")
        
        # Perform second run with language-specific assistant
        final_corrected_text, second_run_mistakes = await perform_second_run(
            first_run_corrected_text, 
            detected_language, 
            first_run_mistakes, 
            number_protector
        )
        
        # Use only the final mistakes from the second run (which includes comprehensive corrections)
        print(f"=== DOCX - FINAL SUMMARY ===")
        print(f"First run mistakes: {len(first_run_mistakes)}")
        print(f"Second run total output: {len(second_run_mistakes)} (should include carried forward + new)")
        print(f"Final output mistakes: {len(second_run_mistakes)}")
        print(f"Language detected: {detected_language}")
        print(f"=== DOCX - END FINAL SUMMARY ===")
        
        # ============ TRACK CHANGES GENERATION - IMPROVED ============
        print("=== DOCX - CREATING TRACK CHANGES ===")
        print(f"Original length: {len(extracted_text)}")
        print(f"Final corrected length: {len(final_corrected_text)}")
        
        # Check if there are actually meaningful changes
        from word_revisions import WordRevisionGenerator
        generator = WordRevisionGenerator()
        
        if not generator.has_meaningful_changes(extracted_text, final_corrected_text):
            print("No meaningful changes detected - creating document without track changes")
            # Create a simple document showing that no changes were needed
            try:
                corrected_docx = create_simple_corrections_docx(extracted_text, final_corrected_text, ["No corrections needed - text already follows CUHK style guidelines."])
            except Exception as e:
                raise HTTPException(
                    status_code=500, 
                    detail=f"Failed to create document with no changes. Error: {str(e)}"
                )
        else:
            print("Creating track changes document with meaningful changes only")
            try:
                corrected_docx = create_tracked_changes_docx(extracted_text, final_corrected_text, second_run_mistakes)
            except Exception as e:
                print(f"Error in track changes generation: {e}")
                # Fallback to simple document
                corrected_docx = create_simple_corrections_docx(extracted_text, final_corrected_text, second_run_mistakes)
        

        
        # Generate filename for the corrected document
        original_name = file.filename.rsplit('.', 1)[0]
        download_filename = f"{original_name}_corrected_two_stage.docx"
        
        # Save the corrected file temporarily
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, download_filename)
        
        try:
            with open(temp_path, 'wb') as f:
                f.write(corrected_docx.getvalue())
        except Exception as e:
            raise HTTPException(
                status_code=500, 
                detail=f"Failed to save corrected document: {str(e)}"
            )
        
        return DocxProofReadResponse(
            original_filename=file.filename,
            mistakes_count=len(second_run_mistakes),
            mistakes=second_run_mistakes,
            status="completed",
            download_filename=download_filename
        )
    
    except Exception as e:
        print(f"DOCX Error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/style-guides")
async def get_style_guides():
    """
    Get list of available style guide files
    """
    try:
        style_guides_dir = "styling_guides"
        if not os.path.exists(style_guides_dir):
            return {"files": []}
        
        files = []
        for filename in os.listdir(style_guides_dir):
            if filename.endswith(('.pdf', '.doc', '.docx', '.zip', '.md', '.json')):
                # Clean up the display name
                display_name = filename.replace('.pdf', '').replace('.doc', '').replace('.docx', '').replace('.zip', '').replace('.md', '').replace('.json', '')
                display_name = display_name.replace('%20', ' ')  # Handle URL encoding
                
                files.append({
                    "filename": filename,
                    "display_name": display_name,
                    "download_url": f"/styling_guides/{filename}"
                })
        
        return {"files": files}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/download-docx/{filename}")
async def download_corrected_docx(filename: str):
    """
    Download the corrected DOCX file
    """
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, filename)
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

def get_or_create_english_assistant():
    """Get or create English-specific assistant for second run"""
    assistant_id = None
    
    # Try to load existing assistant ID
    if os.path.exists(ENGLISH_ASSISTANT_CONFIG_FILE):
        try:
            with open(ENGLISH_ASSISTANT_CONFIG_FILE, 'r') as f:
                config = json.load(f)
                assistant_id = config.get("assistant_id")
        except Exception as e:
            print(f"Error loading English assistant config: {e}")
    
    # Check if the assistant still exists
    if assistant_id:
        try:
            assistant = client.beta.assistants.retrieve(assistant_id)
            print(f"Using existing English assistant: {assistant_id}")
            return assistant
        except Exception as e:
            print(f"Existing English assistant {assistant_id} not found, creating new one: {e}")
            assistant_id = None
    
    # Create new English assistant
    print("Creating new English assistant...")
    model_name = os.getenv("AZURE_OPENAI_MODEL", "gpt-4o")
    english_prompt, _ = load_second_run_prompts()
    
    assistant = client.beta.assistants.create(
        model=model_name,
        name="CUHK_English_Editor",
        instructions=f"""
{english_prompt}

You are performing a comprehensive FINAL proofreading pass. The text has already been through an initial proofreading pass, and you will be given a list of corrections that were already made.

IMPORTANT: Your "mistakes" array should include:
1. ALL corrections from the first pass (with any refinements or improvements)
2. ANY additional corrections you identify
3. This gives users a complete view of all corrections made to their text

Return your response as a JSON object that strictly follows the required schema. Include ALL corrections (both carried forward and newly identified).
        """,
        tools=[{"type": "file_search"}],
        tool_resources={"file_search": {"vector_store_ids": ["vs_GENF8IR41N6uP60Jx9CuLgbs"]}},
        response_format={
            "type": "json_schema",
            "json_schema": {
                "name": "ProofreadingResponse",
                "strict": True,
                "schema": {
                    "type": "object",
                    "properties": {
                        "corrected_text": {
                            "type": "string",
                            "description": "The corrected version of the text"
                        },
                        "mistakes": {
                            "type": "array",
                            "description": "List of mistakes found and how they were corrected with specific CUHK style guide references",
                            "items": {
                                "type": "string"
                            }
                        }
                    },
                    "required": [
                        "corrected_text",
                        "mistakes",
                    ],
                    "additionalProperties": False
                }
            }
        },
        temperature=0.05,
        top_p=0.15
    )
    
    # Save the assistant ID
    try:
        config = {"assistant_id": assistant.id}
        with open(ENGLISH_ASSISTANT_CONFIG_FILE, 'w') as f:
            json.dump(config, f)
        print(f"Saved new English assistant ID: {assistant.id}")
    except Exception as e:
        print(f"Error saving English assistant config: {e}")
    
    return assistant

def get_or_create_chinese_assistant():
    """Get or create Chinese-specific assistant for second run"""
    assistant_id = None
    
    # Try to load existing assistant ID
    if os.path.exists(CHINESE_ASSISTANT_CONFIG_FILE):
        try:
            with open(CHINESE_ASSISTANT_CONFIG_FILE, 'r') as f:
                config = json.load(f)
                assistant_id = config.get("assistant_id")
        except Exception as e:
            print(f"Error loading Chinese assistant config: {e}")
    
    # Check if the assistant still exists
    if assistant_id:
        try:
            assistant = client.beta.assistants.retrieve(assistant_id)
            print(f"Using existing Chinese assistant: {assistant_id}")
            return assistant
        except Exception as e:
            print(f"Existing Chinese assistant {assistant_id} not found, creating new one: {e}")
            assistant_id = None
    
    # Create new Chinese assistant
    print("Creating new Chinese assistant...")
    model_name = os.getenv("AZURE_OPENAI_MODEL", "gpt-4o")
    _, chinese_prompt = load_second_run_prompts()
    
    assistant = client.beta.assistants.create(
        model=model_name,
        name="CUHK_Chinese_Editor",
        instructions=f"""
{chinese_prompt}

您正在進行全面的最終校對。文本已經經過初步校對，您將獲得已進行的修正清單。

重要：您的「mistakes」陣列應包含：
1. 第一輪的所有修正（並可進行任何改進或完善）
2. 您識別出的任何額外修正
3. 這為用戶提供對其文本所有修正的完整視圖

請以JSON格式返回您的回應，嚴格遵循所需的架構。包含所有修正（既有延續的也有新識別的）。
        """,
        tools=[{"type": "file_search"}],
        tool_resources={"file_search": {"vector_store_ids": ["vs_GENF8IR41N6uP60Jx9CuLgbs"]}},
        response_format={
            "type": "json_schema",
            "json_schema": {
                "name": "ProofreadingResponse",
                "strict": True,
                "schema": {
                    "type": "object",
                    "properties": {
                        "corrected_text": {
                            "type": "string",
                            "description": "The corrected version of the text"
                        },
                        "mistakes": {
                            "type": "array",
                            "description": "List of mistakes found and how they were corrected with specific CUHK style guide references",
                            "items": {
                                "type": "string"
                            }
                        }
                    },
                    "required": [
                        "corrected_text",
                        "mistakes",
                    ],
                    "additionalProperties": False
                }
            }
        },
        temperature=0.05,
        top_p=0.15
    )
    
    # Save the assistant ID
    try:
        config = {"assistant_id": assistant.id}
        with open(CHINESE_ASSISTANT_CONFIG_FILE, 'w') as f:
            json.dump(config, f)
        print(f"Saved new Chinese assistant ID: {assistant.id}")
    except Exception as e:
        print(f"Error saving Chinese assistant config: {e}")
    
    return assistant

async def perform_second_run(text: str, language: str, first_run_mistakes: list, number_protector) -> tuple:
    """
    Perform the second run with language-specific assistant
    Returns (corrected_text, mistakes)
    """
    # DISABLED: Skip pattern protection for second run to allow better corrections
    # Pattern protection is only applied in the first run for basic safety
    protected_text = text
    protection_instructions = ""
    print(f"Second run - Pattern protection DISABLED for better correction results ({language} text)")
    
    # Select appropriate assistant based on language
    if language == 'english':
        assistant_instance = get_english_assistant()
        language_note = "This is English text. Apply CUHK English Style Guide rules strictly."
    elif language == 'chinese':
        assistant_instance = get_chinese_assistant()
        language_note = "這是中文文本。請嚴格按照CUHK中文編輯指引執行校對。"
    else:  # mixed
        # For mixed text, use the language that appears more prominently
        chinese_chars = sum(1 for char in text if '\u4e00' <= char <= '\u9fff')
        english_chars = sum(1 for char in text if 'a' <= char.lower() <= 'z')
        
        if chinese_chars > english_chars:
            assistant_instance = get_chinese_assistant()
            language_note = "這是混合語言文本，以中文為主。請按照CUHK中文編輯指引進行校對，同時注意英文部分的規範。"
        else:
            assistant_instance = get_english_assistant()
            language_note = "This is mixed-language text with English predominant. Apply CUHK English Style Guide rules while ensuring Chinese portions are properly formatted."
    
    # Create a new thread for second run
    thread = client.beta.threads.create()
    
    # Build message for second run
    message_content = f"""{language_note}

FIRST RUN CORRECTIONS SUMMARY:
The following {len(first_run_mistakes)} corrections were already made in the first pass:
{chr(10).join([f"• {mistake}" for mistake in first_run_mistakes])}

TEXT TO REVIEW (Second Pass - Already Corrected Once):
{protected_text}

###Please perform a comprehensive FINAL proofreading pass. Review the text thoroughly and provide a COMPLETE list of ALL corrections needed, including:
1. All corrections that were made in the first pass (listed above) - carry these forward with any refinements
2. Any additional corrections you identify that were missed in the first pass
3. Any further refinements according to strict CUHK style guidelines

This is the final review. Your "mistakes" array should include ALL corrections (both from first pass and any new ones) so the user sees the complete list of what was corrected.###
"""
    
    # Note: Protection instructions are disabled for second run to allow better corrections
    
    # Add user message to the thread
    message = client.beta.threads.messages.create(
        thread_id=thread.id,
        role="user",
        content=message_content
    )
    
    # Run the thread
    run = client.beta.threads.runs.create(
        thread_id=thread.id,
        assistant_id=assistant_instance.id
    )
    
    # Wait for completion
    run = wait_for_run_completion(client, thread.id, run.id)
    
    if run.status == 'completed':
        # Get the assistant's response
        messages = client.beta.threads.messages.list(
            thread_id=thread.id
        )
        
        # Extract the assistant's response
        assistant_response = ""
        for message in messages.data:
            if message.role == "assistant":
                try:
                    assistant_response = message.content[0].text.value
                    break
                except (IndexError, AttributeError) as e:
                    print(f"Warning: Error extracting second run message content: {e}")
                    continue
        
        # Validate that we got a response
        if not assistant_response.strip():
            print("Empty response from second run assistant")
            return text, [f"Second run returned empty response"]
        
        print(f"=== SECOND RUN ({language.upper()}) RESPONSE ===")
        print(assistant_response)
        print(f"=== END SECOND RUN ({language.upper()}) RESPONSE ===")
        
        # Parse the second run response with robust error handling
        corrected_text, mistakes = parse_assistant_response(
            clean_response_text(assistant_response), 
            text
        )
        
        print(f"=== SECOND RUN PARSED ({language.upper()}) ===")
        print(f"Corrected text length: {len(corrected_text)}")
        print(f"Total mistakes in final output: {len(mistakes)} (includes carried forward + new)")
        print(f"First run had: {len(first_run_mistakes)} mistakes")
        print(f"=== END SECOND RUN PARSED ({language.upper()}) ===")
        
        
        # Pattern protection was disabled for second run - no restoration needed
        restored_mistakes = mistakes
        print(f"Second run - No pattern restoration needed (protection was disabled for better results)")
        
        return corrected_text, restored_mistakes
    
    else:
        print(f"Second run failed with status: {run.status}")
        return text, [f"Second run failed with status: {run.status}"]
