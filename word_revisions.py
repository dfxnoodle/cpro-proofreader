#!/usr/bin/env python3
"""
Word revisions module for creating DOCX files with track changes.
Handles proper spacing preservation to avoid words sticking together.
"""

import re
import uuid
import xml.etree.ElementTree as ET
from datetime import datetime
from io import BytesIO
from typing import List, Tuple, Union
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import RGBColor
import difflib
import Levenshtein


class WordRevisionGenerator:
    """Generates Word documents with track changes, preserving proper spacing."""
    
    def __init__(self):
        self.revision_id = 0
        self.author = "Proofreader"
        self.date = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    
    def create_document_with_revisions(self, original_text: str, corrected_text: str, mistakes: List[str], citations: List[dict] = None) -> BytesIO:
        """
        Create a DOCX document with track changes showing the differences between original and corrected text.
        
        Args:
            original_text: The original text
            corrected_text: The corrected text
            mistakes: List of mistake descriptions
            citations: List of citation information
            
        Returns:
            BytesIO: The generated DOCX file as bytes
        """
        # Create a new document
        doc = Document()
        
        # Enable track changes in document settings
        self._enable_track_changes(doc)
        
        # Add a title
        title = doc.add_heading("Document with Track Changes", level=1)
        
        # Generate word-level differences
        changes = self._generate_word_diff(original_text, corrected_text)
        
        # Filter out minor changes
        filtered_changes = []
        for change_type, text in changes:
            if not self.should_ignore_change(change_type, text):
                filtered_changes.append((change_type, text))
            else:
                # Convert ignored deletions/insertions to unchanged text
                if change_type == 'delete':
                    # If we're ignoring a deletion, don't show it as deleted
                    filtered_changes.append(('equal', text))
                elif change_type == 'insert':
                    # If we're ignoring an insertion, show it as normal text
                    filtered_changes.append(('equal', text))
                else:
                    filtered_changes.append((change_type, text))
        
        # Create paragraph with proper Word track changes
        para = doc.add_paragraph()
        
        for change_type, text in filtered_changes:
            if change_type == 'equal':
                # Unchanged text - add normally
                para.add_run(text)
            elif change_type == 'delete':
                # Add deleted text using Word's native deletion tracking
                self._add_deletion_to_paragraph(para, text)
            elif change_type == 'insert':
                # Add inserted text using Word's native insertion tracking
                self._add_insertion_to_paragraph(para, text)
        
        # Add corrections summary at the bottom
        if mistakes:
            # Add separator
            doc.add_paragraph("─" * 50)
            
            mistakes_para = doc.add_paragraph()
            mistakes_para.add_run("Corrections Made:").bold = True
            for mistake in mistakes:
                # Use regular paragraph without numbered style to avoid duplication
                mistake_para = doc.add_paragraph(f"• {mistake}")
        
        # Add citations section
        if citations:
            # Add separator if we haven't added one yet
            if not mistakes:
                doc.add_paragraph("─" * 50)
            
            citations_para = doc.add_paragraph()
            citations_para.add_run("Citations and References:").bold = True
            for i, citation in enumerate(citations, 1):
                citation_para = doc.add_paragraph(f"{i}. {citation.get('text', '')}")
                if citation.get('quote'):
                    quote_para = doc.add_paragraph(f"   Quote: \"{citation['quote']}\"")
                    quote_para.italic = True
                if citation.get('file_name'):
                    source_para = doc.add_paragraph(f"   Source: {citation['file_name']}")
                    source_para.italic = True
        
        # Save to BytesIO
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        return doc_bytes
    
    def has_meaningful_changes(self, original_text: str, corrected_text: str) -> bool:
        """
        Check if there are meaningful changes between original and corrected text.
        
        Args:
            original_text: Original text
            corrected_text: Corrected text
            
        Returns:
            True if there are meaningful changes worth showing in track changes
        """
        # Quick check for identical text
        if original_text.strip() == corrected_text.strip():
            return False
        
        # Generate differences and check if any are meaningful
        changes = self._generate_word_diff(original_text, corrected_text)
        
        for change_type, text in changes:
            if change_type in ['delete', 'insert']:
                # If this change wouldn't be ignored, it's meaningful
                if not self.should_ignore_change(change_type, text):
                    return True
        
        return False

    def should_ignore_change(self, change_type: str, text: str) -> bool:
        """
        Determine if a change should be ignored (not shown in track changes).
        Returns True if the change is too minor to show.
        
        Args:
            change_type: Type of change ('delete', 'insert', 'equal')
            text: The text content of the change
            
        Returns:
            True if change should be ignored
        """
        if change_type == 'equal':
            return False  # Never ignore unchanged text
        
        # Ignore pure whitespace changes that are very small
        if not text.strip():
            # Only ignore very minor whitespace changes (1-2 characters)
            return len(text) <= 2
        
        # Ignore single character punctuation or spacing changes that don't add meaning
        if len(text.strip()) == 1 and text.strip() in '.,;: \t\n':
            return True
        
        return False

    def _generate_word_diff(self, original: str, corrected: str) -> List[Tuple[str, str]]:
        """
        Generate differences between original and corrected text using intelligent approach.
        Uses character-level diffing for Chinese text and precise word-level for English.
        
        Args:
            original: Original text
            corrected: Corrected text
            
        Returns:
            List of tuples (change_type, text) where change_type is 'equal', 'delete', or 'insert'
        """
        # Check if text is primarily Chinese
        if self._is_chinese_text(original) or self._is_chinese_text(corrected):
            # Use character-level diff for Chinese text
            return self._generate_character_diff(original, corrected)
        else:
            # Use precise word-level diff for English text
            return self._generate_precise_word_diff(original, corrected)
    
    def _is_chinese_text(self, text: str) -> bool:
        """
        Check if text contains Chinese characters.
        
        Args:
            text: Text to check
            
        Returns:
            True if text contains Chinese characters
        """
        chinese_char_pattern = re.compile(r'[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff]')
        return bool(chinese_char_pattern.search(text))
    
    def _generate_character_diff(self, original: str, corrected: str) -> List[Tuple[str, str]]:
        """
        Generate character-level differences using Levenshtein for precise tracking.
        
        Args:
            original: Original text
            corrected: Corrected text
            
        Returns:
            List of tuples (change_type, text)
        """
        opcodes = Levenshtein.opcodes(original, corrected)
        
        changes = []
        for op, i1, i2, j1, j2 in opcodes:
            if op == 'equal':
                text = original[i1:i2]
                if text:
                    changes.append(('equal', text))
            elif op == 'delete':
                text = original[i1:i2]
                if text:
                    changes.append(('delete', text))
            elif op == 'insert':
                text = corrected[j1:j2]
                if text:
                    changes.append(('insert', text))
            elif op == 'replace':
                deleted_text = original[i1:i2]
                if deleted_text:
                    changes.append(('delete', deleted_text))
                inserted_text = corrected[j1:j2]
                if inserted_text:
                    changes.append(('insert', inserted_text))
        
        return self._merge_consecutive_changes(changes)
    
    def _generate_contextual_word_diff(self, original: str, corrected: str) -> List[Tuple[str, str]]:
        """
        Generate word-level differences with sentence context to avoid bulk replacements.
        
        Args:
            original: Original text
            corrected: Corrected text
            
        Returns:
            List of tuples (change_type, text)
        """
        # Split into sentences first
        original_sentences = self._split_into_sentences(original)
        corrected_sentences = self._split_into_sentences(corrected)
        
        # Use sentence-level matching first
        sentence_matcher = difflib.SequenceMatcher(None, original_sentences, corrected_sentences)
        
        changes = []
        for op, i1, i2, j1, j2 in sentence_matcher.get_opcodes():
            if op == 'equal':
                # Sentences are identical
                for sentence in original_sentences[i1:i2]:
                    changes.append(('equal', sentence))
            elif op == 'replace' and i2 - i1 == 1 and j2 - j1 == 1:
                # Single sentence replacement - use word-level diff
                orig_sentence = original_sentences[i1]
                corr_sentence = corrected_sentences[j1]
                sentence_changes = self._generate_word_level_diff(orig_sentence, corr_sentence)
                changes.extend(sentence_changes)
            else:
                # Multiple sentence changes - treat as delete + insert
                for sentence in original_sentences[i1:i2]:
                    changes.append(('delete', sentence))
                for sentence in corrected_sentences[j1:j2]:
                    changes.append(('insert', sentence))
        
        return self._merge_consecutive_changes(changes)
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """
        Split text into sentences while preserving punctuation and spacing.
        
        Args:
            text: Text to split
            
        Returns:
            List of sentences including their punctuation and following spaces
        """
        if not text:
            return []
        
        # Split on sentence-ending punctuation but keep the punctuation
        sentence_pattern = re.compile(r'([.!?]+\s*)')
        parts = sentence_pattern.split(text)
        
        sentences = []
        current_sentence = ""
        
        for part in parts:
            current_sentence += part
            if sentence_pattern.match(part):
                # This part contains sentence-ending punctuation
                sentences.append(current_sentence)
                current_sentence = ""
        
        # Add any remaining text as the last sentence
        if current_sentence.strip():
            sentences.append(current_sentence)
        
        return [s for s in sentences if s.strip()]
    
    def _generate_word_level_diff(self, original: str, corrected: str) -> List[Tuple[str, str]]:
        """
        Generate word-level differences for a single sentence or short text.
        
        Args:
            original: Original sentence
            corrected: Corrected sentence
            
        Returns:
            List of tuples (change_type, text)
        """
        # Tokenize preserving spaces
        original_tokens = self._tokenize_with_spaces(original)
        corrected_tokens = self._tokenize_with_spaces(corrected)
        
        matcher = difflib.SequenceMatcher(None, original_tokens, corrected_tokens)
        
        changes = []
        for op, i1, i2, j1, j2 in matcher.get_opcodes():
            if op == 'equal':
                text = ''.join(original_tokens[i1:i2])
                if text:
                    changes.append(('equal', text))
            elif op == 'delete':
                text = ''.join(original_tokens[i1:i2])
                if text:
                    changes.append(('delete', text))
            elif op == 'insert':
                text = ''.join(corrected_tokens[j1:j2])
                if text:
                    changes.append(('insert', text))
            elif op == 'replace':
                deleted_text = ''.join(original_tokens[i1:i2])
                if deleted_text:
                    changes.append(('delete', deleted_text))
                inserted_text = ''.join(corrected_tokens[j1:j2])
                if inserted_text:
                    changes.append(('insert', inserted_text))
        
        return changes

    def _tokenize_with_spaces(self, text: str) -> List[str]:
        """
        Tokenize text while preserving spaces and punctuation.
        Each token is either a word, a space, or punctuation.
        
        Args:
            text: Input text to tokenize
            
        Returns:
            List of tokens preserving all spacing and punctuation
        """
        if not text:
            return []
        
        tokens = []
        current_token = ""
        
        for char in text:
            if char.isalnum() or char in "'-":
                # Part of a word
                current_token += char
            else:
                # Not part of a word (space, punctuation, etc.)
                if current_token:
                    tokens.append(current_token)
                    current_token = ""
                if char:  # Don't add empty characters
                    tokens.append(char)
        
        # Add any remaining token
        if current_token:
            tokens.append(current_token)
        
        return tokens
    
    def _enable_track_changes(self, doc):
        """
        Enable track changes in the document settings.
        
        Args:
            doc: The Document object
        """
        # Access the document settings
        settings = doc.settings
        
        # Add track changes setting to the XML
        settings_xml = settings._element
        
        # Create the track changes element if it doesn't exist
        track_changes = settings_xml.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trackRevisions')
        if track_changes is None:
            track_changes_xml = f'''<w:trackRevisions {nsdecls('w')}/>'''
            track_changes_element = parse_xml(track_changes_xml)
            settings_xml.append(track_changes_element)
    
    def _add_deletion_to_paragraph(self, paragraph, text):
        """
        Add deleted text to paragraph using Word's native deletion tracking.
        
        Args:
            paragraph: The paragraph to add to
            text: The deleted text
        """
        self.revision_id += 1
        
        # Create deletion XML with red strikethrough (Word's default for deletions)
        del_xml = f'''
        <w:del {nsdecls('w')} w:id="{self.revision_id}" w:author="{self.author}" w:date="{self.date}">
            <w:r>
                <w:rPr>
                    <w:color w:val="FF0000"/>
                    <w:strike w:val="true"/>
                </w:rPr>
                <w:delText>{self._escape_xml(text)}</w:delText>
            </w:r>
        </w:del>
        '''
        
        # Parse and add to paragraph
        del_element = parse_xml(del_xml)
        paragraph._element.append(del_element)
    
    def _add_insertion_to_paragraph(self, paragraph, text):
        """
        Add inserted text to paragraph using Word's native insertion tracking.
        
        Args:
            paragraph: The paragraph to add to
            text: The inserted text
        """
        self.revision_id += 1
        
        # Create insertion XML with green color for better visibility of corrections
        ins_xml = f'''
        <w:ins {nsdecls('w')} w:id="{self.revision_id}" w:author="{self.author}" w:date="{self.date}">
            <w:r>
                <w:rPr>
                    <w:color w:val="008000"/>
                </w:rPr>
                <w:t xml:space="preserve">{self._escape_xml(text)}</w:t>
            </w:r>
        </w:ins>
        '''
        
        # Parse and add to paragraph
        ins_element = parse_xml(ins_xml)
        paragraph._element.append(ins_element)
    
    def _escape_xml(self, text):
        """
        Escape XML special characters in text.
        
        Args:
            text: Text to escape
            
        Returns:
            Escaped text safe for XML
        """
        if not text:
            return ""
        
        # Replace XML special characters
        text = text.replace("&", "&amp;")
        text = text.replace("<", "&lt;")
        text = text.replace(">", "&gt;")
        text = text.replace('"', "&quot;")
        text = text.replace("'", "&apos;")
        
        return text

    def _create_revision_element(self, text: str, revision_type: str) -> OxmlElement:
        """
        Create a Word revision element with proper XML structure.
        
        Args:
            text: The text content
            revision_type: Type of revision ('insert' or 'delete')
            
        Returns:
            OxmlElement: The revision element
        """
        self.revision_id += 1
        
        if revision_type == 'insert':
            # Create insertion revision - green color for better visibility of corrections
            ins_xml = f'''
            <w:ins {nsdecls('w')} w:id="{self.revision_id}" w:author="{self.author}" w:date="{self.date}">
                <w:r>
                    <w:rPr>
                        <w:color w:val="008000"/>
                    </w:rPr>
                    <w:t>{text}</w:t>
                </w:r>
            </w:ins>
            '''
            return parse_xml(ins_xml)
        
        elif revision_type == 'delete':
            # Create deletion revision - red strikethrough
            del_xml = f'''
            <w:del {nsdecls('w')} w:id="{self.revision_id}" w:author="{self.author}" w:date="{self.date}">
                <w:r>
                    <w:rPr>
                        <w:color w:val="FF0000"/>
                        <w:strike w:val="true"/>
                    </w:rPr>
                    <w:t>{text}</w:t>
                </w:r>
            </w:del>
            '''
            return parse_xml(del_xml)
        
        return None
    
    def _merge_consecutive_changes(self, changes: List[Tuple[str, str]]) -> List[Tuple[str, str]]:
        """
        Merge consecutive changes of the same type to reduce noise.
        
        Args:
            changes: List of (change_type, text) tuples
            
        Returns:
            Optimized list with consecutive same-type changes merged
        """
        if not changes:
            return changes
        
        merged = []
        current_type, current_text = changes[0]
        
        for change_type, text in changes[1:]:
            if change_type == current_type:
                # Same type - merge the text
                current_text += text
            else:
                # Different type - add current and start new
                merged.append((current_type, current_text))
                current_type, current_text = change_type, text
        
        # Add the last accumulated change
        merged.append((current_type, current_text))
        
        return merged

    def _generate_precise_word_diff(self, original: str, corrected: str) -> List[Tuple[str, str]]:
        """
        Generate precise word-level differences for English text.
        This method focuses on making minimal changes rather than bulk replacements.
        
        Args:
            original: Original text
            corrected: Corrected text
            
        Returns:
            List of tuples (change_type, text)
        """
        # Tokenize the text preserving spaces and punctuation
        original_tokens = self._tokenize_with_spaces(original)
        corrected_tokens = self._tokenize_with_spaces(corrected)
        
        # Use difflib for precise sequence matching
        matcher = difflib.SequenceMatcher(None, original_tokens, corrected_tokens)
        
        changes = []
        for op, i1, i2, j1, j2 in matcher.get_opcodes():
            if op == 'equal':
                # Unchanged tokens
                text = ''.join(original_tokens[i1:i2])
                if text:
                    changes.append(('equal', text))
            elif op == 'delete':
                # Deleted tokens
                text = ''.join(original_tokens[i1:i2])
                if text:
                    changes.append(('delete', text))
            elif op == 'insert':
                # Inserted tokens
                text = ''.join(corrected_tokens[j1:j2])
                if text:
                    changes.append(('insert', text))
            elif op == 'replace':
                # Replace tokens - show as delete old + insert new
                deleted_text = ''.join(original_tokens[i1:i2])
                if deleted_text:
                    changes.append(('delete', deleted_text))
                inserted_text = ''.join(corrected_tokens[j1:j2])
                if inserted_text:
                    changes.append(('insert', inserted_text))
        
        # Optimize consecutive changes
        return self._merge_consecutive_changes(changes)

def create_word_track_changes_docx(original_text: str, corrected_text: str, mistakes: List[str], citations: List[dict] = None) -> BytesIO:
    """
    Create a Word document with track changes showing differences between original and corrected text.
    This is the main function called by the FastAPI application.
    
    Args:
        original_text: The original text
        corrected_text: The corrected text  
        mistakes: List of mistake descriptions
        citations: List of citation information
        
    Returns:
        BytesIO: The generated DOCX file as bytes
    """
    generator = WordRevisionGenerator()
    return generator.create_document_with_revisions(original_text, corrected_text, mistakes, citations)


# Alternative implementation with more sophisticated tracking
class AdvancedWordRevisionGenerator(WordRevisionGenerator):
    """Advanced version with better character-level tracking for complex cases."""
    
    def create_document_with_revisions(self, original_text: str, corrected_text: str, mistakes: List[str], citations: List[dict] = None) -> BytesIO:
        """
        Create a DOCX document with advanced track changes and better spacing handling.
        """
        # Use the parent method but with enhanced diff generation
        return super().create_document_with_revisions(original_text, corrected_text, mistakes, citations)
    
    def _generate_word_diff(self, original: str, corrected: str) -> List[Tuple[str, str]]:
        """
        Enhanced word diff that handles edge cases better.
        """
        # Handle the case where text has no spaces (all stuck together)
        if ' ' not in original and ' ' in corrected:
            return self._handle_spacing_insertion(original, corrected)
        
        # Handle normal case with existing implementation
        return super()._generate_word_diff(original, corrected)
    
    def _handle_spacing_insertion(self, original: str, corrected: str) -> List[Tuple[str, str]]:
        """
        Handle the special case where spaces need to be inserted into stuck-together text.
        
        Args:
            original: Text without proper spacing
            corrected: Text with proper spacing
            
        Returns:
            List of change tuples
        """
        changes = []
        
        # Use character-level diffing for this case
        original_chars = list(original)
        corrected_chars = list(corrected)
        
        matcher = difflib.SequenceMatcher(None, original_chars, corrected_chars)
        
        for op, i1, i2, j1, j2 in matcher.get_opcodes():
            if op == 'equal':
                text = ''.join(original_chars[i1:i2])
                if text:
                    changes.append(('equal', text))
            elif op == 'delete':
                text = ''.join(original_chars[i1:i2])
                if text:
                    changes.append(('delete', text))
            elif op == 'insert':
                text = ''.join(corrected_chars[j1:j2])
                if text:
                    changes.append(('insert', text))
            elif op == 'replace':
                # For replace, check if we're just adding spaces
                deleted = ''.join(original_chars[i1:i2])
                inserted = ''.join(corrected_chars[j1:j2])
                
                # If the inserted text contains the deleted text plus spaces, handle specially
                if deleted in inserted and ' ' in inserted:
                    # Mark original as deleted and new as inserted
                    changes.append(('delete', deleted))
                    changes.append(('insert', inserted))
                else:
                    # Regular replace
                    if deleted:
                        changes.append(('delete', deleted))
                    if inserted:
                        changes.append(('insert', inserted))
        
        return changes


# For testing with the advanced generator
def create_word_track_changes_docx_advanced(original_text: str, corrected_text: str, mistakes: List[str], citations: List[dict] = None) -> BytesIO:
    """
    Create a Word document using the advanced revision generator.
    """
    generator = AdvancedWordRevisionGenerator()
    return generator.create_document_with_revisions(original_text, corrected_text, mistakes, citations)
